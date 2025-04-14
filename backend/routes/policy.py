from flask import Blueprint, request, jsonify, current_app
from flask_jwt_extended import jwt_required, current_user
from werkzeug.utils import secure_filename
from models import db, PolicyDocument, PolicyChunk
from datetime import datetime, timezone

import os

# Local embedding model disabled; relying on external service or stub

import requests
from config import Config

def embed_text(text):
    """
    Generate embedding using external Ollama server.
    """
    try:
        response = requests.post(
            f"{Config.OLLAMA_API_URL}/embeddings",
            json={
                "model": "nomic-embed-text",  # Change to your Ollama embedding model name if different
                "prompt": text
            },
            timeout=30
        )
        response.raise_for_status()
        data = response.json()
        return data.get("embedding", [])
    except Exception as e:
        # Log error and return dummy embedding to avoid crash
        print(f"Ollama embedding error: {e}")
        return [0.0] * 768

policy_bp = Blueprint('policy', __name__, url_prefix='/api/policies')

@policy_bp.route('/upload', methods=['POST'])
@jwt_required()
def upload_policy():
    """
    Upload a policy document file, extract text, save document and chunks.
    """
    if 'file' not in request.files:
        return jsonify({'error': 'No file part in request'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    filename = secure_filename(file.filename)
    file_type = os.path.splitext(filename)[1].lower().strip('.')

    try:
        # Read file content
        file_bytes = file.read()

        # Save file to disk as well as to DB
        UPLOAD_FOLDER = 'inputs/policy_uploads'
        os.makedirs(UPLOAD_FOLDER, exist_ok=True)
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        disk_filename = f"{timestamp}_{filename}"
        save_path = os.path.join(UPLOAD_FOLDER, disk_filename)
        with open(save_path, "wb") as f:
            f.write(file_bytes)
        text_content = ""

        # Text extraction based on file type
        if file_type == 'txt':
            text_content = file_bytes.decode('utf-8', errors='ignore')

        elif file_type == 'pdf':
            try:
                import io
                import pdfplumber
                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    pages = [page.extract_text() or "" for page in pdf.pages]
                    text_content = "\n\n".join(pages)
            except Exception as e:
                current_app.logger.error(f"PDF extraction failed: {e}", exc_info=True)
                text_content = "[Error extracting text from PDF.]"

        elif file_type in ['docx', 'doc']:
            try:
                import io
                from docx import Document
                doc = Document(io.BytesIO(file_bytes))
                paragraphs = [p.text for p in doc.paragraphs]
                text_content = "\n\n".join(paragraphs)
            except Exception as e:
                current_app.logger.error(f"DOCX extraction failed: {e}", exc_info=True)
                text_content = "[Error extracting text from DOCX.]"

        else:
            text_content = "[Unsupported file type for text extraction.]"

        # Save PolicyDocument with status "Pending"
        new_doc = PolicyDocument(
            filename=filename,
            file_type=file_type,
            uploaded_at=datetime.now(timezone.utc),
            uploader_id=current_user.id,
            content=text_content,
            file_data=file_bytes,
            status="Pending",
            chunk_count=0,
            error_message=None,
            file_path=save_path
        )
        db.session.add(new_doc)
        db.session.flush()  # Get new_doc.id before commit

        # Chunking: simple split by paragraphs
        try:
            paragraphs = [p.strip() for p in text_content.split('\n\n') if p.strip()]
            chunk_count = 0
            for para in paragraphs:
                embedding = embed_text(para)
                chunk = PolicyChunk(
                    document_id=new_doc.id,
                    chunk_text=para,
                    embedding=embedding,
                    created_at=datetime.now(timezone.utc)
                )
                db.session.add(chunk)
                chunk_count += 1
            new_doc.chunk_count = chunk_count
            new_doc.status = "Indexed"
            new_doc.error_message = None
            db.session.commit()
        except Exception as chunk_err:
            db.session.rollback()
            new_doc.status = "Error"
            new_doc.error_message = f"Chunking/embedding error: {chunk_err}"
            db.session.add(new_doc)
            db.session.commit()
            current_app.logger.error(f"Chunking/embedding error: {chunk_err}", exc_info=True)
            return jsonify({'error': 'Failed to process document chunks', 'details': str(chunk_err)}), 500

        # Ingest into FAISS index for vector search
        try:
            from utils.llamaindex_faiss import ingest_policy_document
            ingest_policy_document(new_doc.id, text_content)
        except Exception as faiss_err:
            new_doc.status = "Error"
            new_doc.error_message = f"FAISS ingestion failed: {faiss_err}"
            db.session.add(new_doc)
            db.session.commit()
            current_app.logger.error(f"FAISS ingestion failed: {faiss_err}", exc_info=True)
            return jsonify({'error': 'FAISS ingestion failed', 'details': str(faiss_err)}), 500

        return jsonify({'message': 'Policy uploaded successfully', 'policy_id': new_doc.id}), 201

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error uploading policy: {str(e)}", exc_info=True)
        return jsonify({'error': 'Failed to upload policy document'}), 500

@policy_bp.route('/', methods=['GET'])
@jwt_required()
def list_policies():
    """
    List all uploaded policy documents.
    """
    try:
        policies = PolicyDocument.query.filter(
            PolicyDocument.file_type.in_(['txt', 'pdf', 'docx', 'doc'])
        ).order_by(PolicyDocument.uploaded_at.desc()).all()
        return jsonify([p.to_dict() for p in policies]), 200
    except Exception as e:
        current_app.logger.error(f"Error listing policies: {str(e)}", exc_info=True)
        return jsonify({'error': 'Failed to fetch policies'}), 500

@policy_bp.route('/<int:policy_id>/view', methods=['GET'])
@jwt_required()
def view_policy(policy_id):
    """
    Return the full extracted text of a policy document.
    """
    try:
        policy = PolicyDocument.query.get_or_404(policy_id)
        return (
            policy.content,
            200,
            {'Content-Type': 'text/plain; charset=utf-8'}
        )
    except Exception as e:
        current_app.logger.error(f"Error fetching policy content: {str(e)}", exc_info=True)
        return jsonify({'error': 'Failed to fetch policy content'}), 500

@policy_bp.route('/<int:policy_id>/file', methods=['GET'])
@jwt_required()
def download_policy_file(policy_id):
    """
    Return the original uploaded file (PDF, DOCX, etc).
    """
    try:
        policy = PolicyDocument.query.get_or_404(policy_id)
        if not policy.file_data:
            return jsonify({'error': 'No original file data available'}), 404

        mime_type = 'application/octet-stream'
        if policy.file_type == 'pdf':
            mime_type = 'application/pdf'
        elif policy.file_type in ['doc', 'docx']:
            mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif policy.file_type == 'txt':
            mime_type = 'text/plain; charset=utf-8'

        return (
            policy.file_data,
            200,
            {
                'Content-Type': mime_type,
                'Content-Disposition': f'inline; filename="{policy.filename}"'
            }
        )
    except Exception as e:
        current_app.logger.error(f"Error fetching original policy file: {str(e)}", exc_info=True)
        return jsonify({'error': 'Failed to fetch original policy file'}), 500

@policy_bp.route('/<int:policy_id>', methods=['DELETE'])
@jwt_required()
def delete_policy(policy_id):
    """
    Delete a policy document and its chunks.
    """
    try:
        policy = PolicyDocument.query.get_or_404(policy_id)
        db.session.delete(policy)
        db.session.commit()
        return jsonify({'message': 'Policy deleted successfully'}), 200
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error deleting policy: {str(e)}", exc_info=True)
        return jsonify({'error': 'Failed to delete policy'}), 500


@policy_bp.route('/search', methods=['POST'])
@jwt_required()
def search_policies():
    """
    Vector similarity search over policy chunks using LlamaIndex + FAISS.
    """
    from utils.llamaindex_faiss import search_policy_chunks

    data = request.get_json()
    query_text = data.get('query', '')
    top_k = data.get('top_k', 5)

    if not query_text:
        return jsonify({"error": "Missing 'query' in request body."}), 400

    results = search_policy_chunks(query_text, top_k)
    return jsonify({"results": results}), 200

@policy_bp.route('/reindex', methods=['POST'])
@jwt_required()
def reindex_policies():
    """
    Re-index all policy documents: re-chunk, re-embed, and update FAISS index.
    """
    from utils.llamaindex_faiss import initialize_faiss_index, ingest_policy_document

    try:
        # Reset FAISS index and metadata
        initialize_faiss_index()
        docs = PolicyDocument.query.all()
        reindexed = 0
        for doc in docs:
            try:
                # Remove old chunks
                PolicyChunk.query.filter_by(document_id=doc.id).delete()
                db.session.flush()
                # Re-chunk and embed
                paragraphs = [p.strip() for p in doc.content.split('\n\n') if p.strip()]
                chunk_count = 0
                for para in paragraphs:
                    embedding = embed_text(para)
                    chunk = PolicyChunk(
                        document_id=doc.id,
                        chunk_text=para,
                        embedding=embedding,
                        created_at=datetime.now(timezone.utc)
                    )
                    db.session.add(chunk)
                    chunk_count += 1
                doc.chunk_count = chunk_count
                doc.status = "Indexed"
                doc.error_message = None
                db.session.commit()
                # Re-ingest into FAISS
                ingest_policy_document(doc.id, doc.content)
                reindexed += 1
            except Exception as err:
                db.session.rollback()
                doc.status = "Error"
                doc.error_message = f"Reindex error: {err}"
                db.session.add(doc)
                db.session.commit()
                current_app.logger.error(f"Reindex error for doc {doc.id}: {err}", exc_info=True)
        return jsonify({"message": f"Re-indexed {reindexed} documents."}), 200
    except Exception as e:
        current_app.logger.error(f"Error during reindex: {e}", exc_info=True)
        return jsonify({"error": "Failed to re-index documents", "details": str(e)}), 500

# The previous pgvector-based search endpoint has been removed.
